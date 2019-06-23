#!/usr/bin/env python3
# -*- coding: utf-8 -*-

'''
Time    :2019/3/14 下午 7:54
Auther  :Mardan
Web     :https://github.com/ka1i
Version :python 3.7
'''

import os
import re
import sys
import time
import timeit
import zipfile
import requests
import configparser
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup
from googletrans import Translator
from docx.enum.text import WD_ALIGN_PARAGRAPH


class GoogleTranslate:
    def __init__(self, src='en', target='zh-cn'):
        self.translator = Translator(service_urls=['translate.google.cn'])
        self.src = src
        self.trg = target

    def translate(self, sentence, max_try=9, count=0):
        # text is string
        if count >= 3:
            return ''
        try:
            translation = self.translator.translate(sentence, dest='zh-cn', src='en')
        except:
            # time.sleep(5)
            translation = self.translate(sentence, time_start , max_try, count=count + 1)
        finally:
            time.sleep(0.1)

        return translation

    def translate_batch(self, sentences, time_start, max_try=10, count=0):
        # sentences is list
        results = []
        if count >= max_try:
            return ['' for _ in range(len(sentences))]
        try:
            translations = self.translator.translate(sentences, dest=self.trg, src=self.src)
            for translation in translations:
                try:
                    translation = translation.text
                except:
                    print('translate error...')
                    translation = ''
                results.append(translation)
        except:
            time.sleep(2.5)
            print('\n\033[1A\033[K' + '翻译重试次数' + str(count + 1)+ ' 耗费时间: ' + str(time_start) + "--->" + str(time.asctime(time.localtime(time.time()))) + ' 秒\033[0m' , end='')
            if count == 5:
                sys.exit(0)
            results = self.translate_batch(sentences, time_start=time.asctime(time.localtime(time.time())), max_try=10, count=count + 1)
        finally:
            # time.sleep(0.5)
            pass

        return results

def getHTMLText(url):
    headers = {
        'User-Agent':
        ('Mozilla/5.0 (Macintosh;Intel Mac OS X 10_12_3) \
            AppleWebKit/537.36 (KHTML, like Gecko) \
            Chrome/56.0.2924.87 Safari/537.36')
    }
    r = requests.get(url, timeout=30, headers=headers)
    try:
        r.raise_for_status()
        r.encoding = r.apparent_encoding
        return r.text
    except Exception as exc:
        return exc

def prepare_cwec_file():
    print("查询CWE常见的弱点最新存档")
    html = getHTMLText('http://cwe.mitre.org/data/archive.html')
    cwe_soup = BeautifulSoup(html, "html.parser")
    cwe_version = cwe_soup.h2.string.split()[-1]
    print("最新存档版本为" + cwe_version)
    cwe_filename_zip = "cwec_v" + cwe_version + ".xml.zip"
    print("正在下载最新版本存档" + cwe_filename_zip)
    cwe_file_context = requests.get(
        'http://cwe.mitre.org/data/xml/' + cwe_filename_zip,
        stream=True
    )
    chunk_size = 1024
    context_size = int(cwe_file_context.headers['content-length'])
    if not os.path.exists('temp'):
        os.mkdir('temp')
    cwe_file = open('temp/' + cwe_filename_zip, 'wb')
    count = 0
    last_per = 0
    after_per = 0
    time_srt = time.time()
    for chunk in cwe_file_context.iter_content(chunk_size=chunk_size):
        if chunk:
            count += len(chunk)
            cwe_file.write(chunk)
            last_per = (count / context_size) * 100
            speed = (
                (
                    (
                        (last_per - after_per) / 100) * context_size) / (time.time() - time_srt) / 1024
            )
            print(
                '\n\033[1A\033[K' + '下载进度: ' +
                '{:.2f}'.format(last_per) + '%' +
                '\t下载速度: ' + '{:.0f}'.format(speed) + 'KB/S'
                '\t剩余时间: ' +
                '{:.2f}'.format(((100 - last_per) / 100) * context_size / (speed * 1024)) +
                '秒' + '\033[0m', end='')
            time_srt = time.time()
            after_per = last_per
    print('\n' + cwe_filename_zip + "下载完成")
    cwe_file.close()
    print('开始解压压缩文件: ' + cwe_filename_zip)
    cwe_file = zipfile.ZipFile('temp/' + cwe_filename_zip, 'r')
    for file in cwe_file.namelist():
        cwe_file.extract(file, "temp")
        print("压缩文件: " + cwe_filename_zip + " 解压成功 ---> ./temp/" + file)

    return file

def Analysis_xml(cwe_filename, start_index, cwe_config):
    xml_context = open('temp/' + cwe_filename, encoding='utf-8').read()
    print("开始 xml 解析文件")
    soup = BeautifulSoup(xml_context, 'xml')
    print("解析 xml 文件成功")
    if not os.path.exists('temp'):
        os.mkdir('temp')
    Weakness_id_index = []
    Weakness_List = []
    for weakness in soup.find_all('Weakness'):
        Weakness_id_index.append(int(weakness.get('ID')))
    Weakness_id_index.sort()
    for cwe_index in range(start_index, len(Weakness_id_index)):
        for weakness in soup.find_all('Weakness'):
            if Weakness_id_index[cwe_index] == int(weakness.get('ID')):
                
                # 获取CWE List Weakness ID
                ID = str(Weakness_id_index[cwe_index])
                print("ID号为: " + ID)
                print("index id号为: " + str(cwe_index))

                # 获取CWE List Weakness Previous Entry Name Date
                Previous_Entry_Name = weakness.Previous_Entry_Name
                if Previous_Entry_Name:
                    Previous_Entry_Name = Previous_Entry_Name.get('Date')
                    print("提交日期: " + Previous_Entry_Name)
                else:
                    Previous_Entry_Name = "I`m don`t know"
                    print("提交日期: " + Previous_Entry_Name)
                
                # 获取CWE List Weakness Modification Date

                if weakness.find_all('Modification'):
                    for modificationdate in weakness.find_all('Modification'):
                        pass
                    Modification_Date = modificationdate.Modification_Date.string
                    print("修改日期: " + Modification_Date)
                else:
                    Modification_Date = "I`m don`t know"
                    print("修改日期: " + Modification_Date)

                # 获取CWE List Weakness Name
                WeaknessName = weakness.get('Name')
                Weakness_List.append(WeaknessName)
                print("CWE弱点名称    : " + WeaknessName)

                # 获取CWE List Weakness Abstraction
                Abstraction = weakness.get('Abstraction')
                print("CWE弱点抽象化  : " + Abstraction)

                # 获取CWE List Weakness Structure
                Structure = weakness.get('Structure')
                print("CWE弱点结构体  : " + Structure)

                # 获取CWE List Weakness Status
                Status = weakness.get('Status')
                print("CWE弱点状态    : " + Status)

                # 获取CWE List Weakness Applicable_Platforms
                WeaknessApplicable_Platforms = weakness.Applicable_Platforms
                Applicable_Platforms = ""
                print("CWE 应用平台 : ", end='')
                if Applicable_Platforms:
                    WeaknessAP_list = (((Applicable_Platforms).next_element).next_element).attrs
                    Applicable_Platforms = str((list((WeaknessAP_list).values()))[0] + "\t" + (list(WeaknessAP_list))[1] + ": " + (list((WeaknessAP_list).values()))[1])
                    print(Applicable_Platforms)
                else:
                    Applicable_Platforms = "没有应用平台"
                    print(Applicable_Platforms)


                # 获取CWE List Weakness Description
                WeaknessDescription = ' '.join((weakness.Description.string).split())
                Weakness_List.append(WeaknessDescription)
                print("CWE弱点简单描述 : " +  WeaknessDescription)

                # 获取CWE List Weakness Extended_Description
                WeaknessExtended_Description = weakness.Extended_Description
                print("CWE弱点详细描述 : ", end='')
                Extended_Description = ""
                if WeaknessExtended_Description:
                    Extended_Description = WeaknessExtended_Description.text
                    print(Extended_Description)
                else:
                    Extended_Description = "没有详细描述"
                    print(Extended_Description)
                Weakness_List.append(Extended_Description)

                # 获取CWE List Weakness Background Detail
                WeaknessBackground_Detail = weakness.Background_Detail
                print("CWE弱点问题背景 : ", end='')
                Background_Detail = ""
                if WeaknessBackground_Detail:
                    Background_Detail = ' '.join((WeaknessBackground_Detail.text).split())
                    print()
                else:
                    Background_Detail = "没有问题背景"
                    print(Background_Detail)
                Weakness_List.append(Background_Detail)

                # 获取CWE List Weakness Notes
                WeaknessNotes = weakness.Notes
                print("CWE弱点笔记 : ", end='')
                Notes = ""
                if WeaknessNotes:
                    Notes = ' '.join((WeaknessNotes.Note.text).split())
                    print(Notes)
                else:
                    Notes = "没有笔记"
                    print(Notes)
                Weakness_List.append(Notes)

                print("\n\033[40m *** ---> 翻译" + "CWE-" + ID + "开始 <--- *** \033[0m\n")
                google_tr = GoogleTranslate()
                cwe_tr_results = google_tr.translate_batch(sentences=Weakness_List, time_start=time.asctime(time.localtime(time.time())))
                for index in range(0, len(cwe_tr_results)):
                    print("\n原文: \n" + Weakness_List[index] + "\n翻译: \n" + cwe_tr_results[index])
                print("\n\033[40m *** ---> 翻译" + "CWE-" + ID + "结束 <--- *** \033[0m\n")
                print("\n\033[40m *** ---> 将" + "CWE-" + ID + "内容写入文档开始 <--- *** \033[0m\n")
                
                cwe_config.set("base", "index_ID", str(cwe_index + 1))
                cwe_config.set("base", "id", str(ID))
                cwe_config.write(open('cwe.config', 'w'))

                document_format(ID, Previous_Entry_Name, Modification_Date, cwe_tr_results[0], Weakness_List[0], Abstraction, Structure, Status, Applicable_Platforms, cwe_tr_results[1], Weakness_List[1], cwe_tr_results[2], Weakness_List[2], cwe_tr_results[3], cwe_tr_results[4])
                print("\n\033[40m *** ---> 将" + "CWE-" + ID + "内容写入文档结束 <--- *** \033[0m\n")
                print("\n\033[34m *** ---> I`m 分割线 <--- *** \033[0m\n")
                Weakness_List.clear()

def write_header(document):
    document.add_heading('CWE翻译计划', 0)
    p = document.add_paragraph('CWE翻译计划 v1.0')
    p.add_run(' by')
    p.add_run(' UESTC').italic = True
    p.add_run(' 418').bold = True
    # 主标题
    document.add_heading('参与贡献人员：', level=1)
    # 副标题
    document.add_paragraph(
        'Mardan & Szg(shang Cr7-joker) & LJIJCJ & Michael Tan',
        style='Intense Quote'
    )
    document.add_page_break()
    document.save('./document/CWE翻译计划.docx')

def document_format(ID, Previous_Entry_Name, Modification_Date, Name_zh, Name_en, Abstraction, Structure, Status, Applicable_Platforms, Description_zh, Description_en, Extended_Description_zh, Extended_Description_en, Background_Detail, Notes):
    document = Document('./document/CWE翻译计划.docx')
    document.add_heading('Weakness ID: ' + ID, level=0)

    one = document.add_paragraph()
    one.add_run('提交日期 ' + Previous_Entry_Name).font.size = Pt(6)
    one.add_run('---> 修改日期 ' + Modification_Date).font.size = Pt(6)
    one.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    one = document.add_paragraph(style='List Bullet')

    one.add_run('Weakness Name:').bold = True
    one.add_run("\n\ten: --> " + Name_en).italic = True
    one.add_run("\n\tzh: -->" + Name_zh).italic = True

    one = document.add_paragraph(style='List Bullet')
    one.add_run('Abstraction: ').bold = True
    one.add_run(Abstraction).italic = True
    one.add_run('\tStructure: ').bold = True
    one.add_run(Structure).italic = True
    one.add_run('\tStatus: ').bold = True
    one.add_run(Status).italic = True
    one = document.add_paragraph(style='List Bullet')
    one.add_run('Applicable_Platforms: ').bold = True
    one.add_run(Applicable_Platforms).italic = True
    one.add_run('\n')   

    document.add_heading('简单描述:', level=3)
    one = document.add_paragraph()
    one.add_run("\t" + Description_zh, 'Body Text Char')
    one.add_run('\n')   

    document.add_heading('Description:', level=3)
    one = document.add_paragraph()
    one.add_run("\t" + Description_en, 'Body Text Char')
    one.add_run('\n')

    document.add_heading('详细描述:', level=3)
    one = document.add_paragraph()
    one.add_run("\t" + Extended_Description_zh, 'Body Text Char')
    one.add_run('\n')   

    document.add_heading('Extended Description:', level=3)
    one = document.add_paragraph()
    one.add_run("\t" + Extended_Description_en, 'Body Text Char')
    one.add_run('\n') 

    document.add_heading('问题背景 (Background Detail):', level=3)
    one = document.add_paragraph()
    one.add_run("\t" + Background_Detail, 'Body Text Char')
    one.add_run('\n')   

    document.add_heading('笔记 (Notes):', level=3)
    one = document.add_paragraph()
    one.add_run("\t" + Notes, 'Body Text Char')
    one.add_run('\n')   

    document.add_page_break()
    document.save('./document/CWE翻译计划.docx')


if __name__ == '__main__':
    start_index = 0
    cwe_config = configparser.ConfigParser()
    time_start = time.time()
    print("正在查询最新版本存档")
    official_cwe_xmlfile = prepare_cwec_file()
    if not os.path.exists('./document'):
        os.mkdir('./document')

    cwe_config.read('cwe.config', encoding='utf8')
    start_index = cwe_config.getint("base", "index_ID")
    w_id = cwe_config.getint("base", "id")
    
    print("开始进行翻译并写入文档")
    if start_index == 0:
        document = Document()
        write_header(document)
    Analysis_xml(official_cwe_xmlfile, start_index, cwe_config)
    print("翻译时间: ", str(time.time() - time_start) + "秒")
   
